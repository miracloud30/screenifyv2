"""
CV Processor Lambda for the AI Job Application Screener.

Triggered by S3 ObjectCreated on uploads/. Extracts CV text, enriches it
Scores it against job requirements using Bedrock
Claude Haiku 4.5, routes the application, persists results, and triggers
the Notifier.

Text extraction order (cheapest and fastest first):
    1. PyPDF2         - text-based PDFs, in-process, multi-page safe
    2. python-docx    - DOCX files
    3. Async Textract - scanned / image PDFs, S3-based, multi-page safe

The old synchronous-bytes Textract path and the utf-8 decode path were
deliberately removed: sync Textract only accepts single-page documents,
and decoding raw PDF bytes as utf-8 never works.
"""

import json
import os
import io
import re
import time
from datetime import datetime, timezone
from decimal import Decimal

import boto3

# Region note: Textract must run in the same region as the uploads bucket
# (eu-north-1). Cross-region Textract->S3 access fails with
# InvalidS3ObjectException, which was one of the original bugs.
REGION = os.environ.get("TEXTRACT_REGION", "eu-north-1")

s3_client = boto3.client("s3", region_name=REGION)
textract_client = boto3.client("textract", region_name=REGION)
bedrock_client = boto3.client(
    "bedrock-runtime", region_name=os.environ.get("BEDROCK_REGION", "us-east-1")
)
dynamodb = boto3.resource("dynamodb", region_name=REGION)
lambda_client = boto3.client("lambda", region_name=REGION)

APPLICATIONS_TABLE = os.environ["APPLICATIONS_TABLE"]
JOBS_TABLE = os.environ["JOBS_TABLE"]
PROCESSED_BUCKET = os.environ["CV_PROCESSED_BUCKET"]
BEDROCK_MODEL = os.environ.get(
    "BEDROCK_MODEL", "us.anthropic.claude-haiku-4-5-20251001-v1:0"
)
NOTIFIER_FUNCTION = "ai-screener-notifier"

# DynamoDB reserved words that appear as attribute names in this app.
# Every reference to these in an UpdateExpression MUST use a placeholder.
RESERVED = {"status", "score", "confidence", "summary", "route"}


def _now():
    return datetime.now(timezone.utc).isoformat()


def _build_update(fields):
    """
    Build a reserved-word-safe UpdateExpression from a {name: value} dict.
    Returns (expression, names, values). Every attribute name is aliased
    with '#', so reserved words like score/status/summary never break.
    """
    sets, names, values = [], {}, {}
    for i, (name, val) in enumerate(fields.items()):
        n, v = f"#n{i}", f":v{i}"
        sets.append(f"{n} = {v}")
        names[n] = name
        values[v] = val
    return "SET " + ", ".join(sets), names, values


def lambda_handler(event, context):
    print(f"Received event: {json.dumps(event)}")
    try:
        record = event["Records"][0]
        bucket = record["s3"]["bucket"]["name"]
        key = record["s3"]["object"]["key"]
        print(f"Processing: s3://{bucket}/{key}")

        tenant_id, job_id, applicant_id = parse_key(key)
        print(f"Tenant: {tenant_id}  Job: {job_id}  Applicant: {applicant_id}")
        update_status(applicant_id, "processing")

        cv_text = extract_document_text(bucket, key)
        print(f"Extracted {len(cv_text)} characters")

        requirements = get_job_requirements(tenant_id, job_id)
        print(f"Loaded requirements for {tenant_id}/{job_id}")

        score_result = score_cv(cv_text, requirements)
        print(f"Scoring result: {json.dumps(score_result)}")

        ai_suggestion = determine_route(score_result, requirements)
        print(f"AI suggestion: {ai_suggestion}")

        # The AI triages; a human decides. We persist the score + the AI's
        # suggested action, but the real status becomes 'needs_review' (or
        # 'flagged' when confidence is low). No candidate email fires here —
        # candidates are only emailed when a recruiter clicks a decision.
        final_status = "flagged" if ai_suggestion == "flagged" else "needs_review"
        persist_results(applicant_id, score_result, ai_suggestion, final_status)
        move_to_processed(bucket, key, applicant_id)
        update_status(applicant_id, final_status)

        # Recruiter-facing alert only (never the candidate) for flagged cases.
        if final_status == "flagged":
            trigger_recruiter_alert(applicant_id, score_result)

        return {
            "statusCode": 200,
            "body": json.dumps(
                {"applicant_id": applicant_id, "ai_suggestion": ai_suggestion,
                 "status": final_status, "score": score_result["score"]}
            ),
        }
    except Exception as e:
        import traceback
        print(f"Error processing CV: {e}")
        print(traceback.format_exc())
        # Best-effort: mark the record so it doesn't sit in 'processing'.
        try:
            _, _, aid = parse_key(event["Records"][0]["s3"]["object"]["key"])
            update_status(aid, "error")
        except Exception:
            pass
        return {"statusCode": 500, "body": str(e)}


def parse_key(key):
    """
    Keys are 'uploads/<tenant_id>/<job_id>/<applicant_id>-<filename>'.
    applicant_id is 'app_<ts>_<uuid8>' (no hyphens), so splitting the last
    path segment on the first hyphen isolates it even if the filename has
    hyphens. Returns (tenant_id, job_id, applicant_id).
    """
    parts = key.split("/")
    # parts: ['uploads', tenant, job, '<applicant_id>-<filename>']
    if len(parts) >= 4 and parts[0] == "uploads":
        tenant_id, job_id = parts[1], parts[2]
        last = parts[-1]
    else:
        # Fallback for legacy/manual keys without tenant/job scoping.
        tenant_id, job_id, last = "default", "default", parts[-1]
    m = re.match(r"^(app_\d+_[0-9a-fA-F]+)-", last)
    applicant_id = m.group(1) if m else last.rsplit(".", 1)[0]
    return tenant_id, job_id, applicant_id


def extract_document_text(bucket, key):
    methods_tried = []
    obj = s3_client.get_object(Bucket=bucket, Key=key)
    document_bytes = obj["Body"].read()
    ext = key.lower().rsplit(".", 1)[-1] if "." in key else ""

    # 1. PyPDF2 for text-based PDFs.
    if ext == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(document_bytes))
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
            if text.strip():
                print(f"PyPDF2 succeeded: {len(text)} characters")
                return text
            methods_tried.append("PyPDF2: empty (likely scanned PDF)")
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
            methods_tried.append(f"PyPDF2: {e}")

    # 2. python-docx for DOCX.
    if ext == "docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(document_bytes))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" ".join(c.text for c in row.cells))
            text = "\n".join(parts)
            if text.strip():
                print(f"python-docx succeeded: {len(text)} characters")
                return text
            methods_tried.append("python-docx: empty")
        except Exception as e:
            print(f"python-docx failed: {e}")
            methods_tried.append(f"python-docx: {e}")

    # 3. Async Textract for scanned / multi-page PDFs (S3-based).
    try:
        text = extract_with_async_textract(bucket, key)
        if text.strip():
            print(f"Async Textract succeeded: {len(text)} characters")
            return text
        methods_tried.append("Async Textract: empty")
    except Exception as e:
        print(f"Async Textract failed: {e}")
        methods_tried.append(f"Async Textract: {e}")

    raise Exception("All extraction methods failed: " + "; ".join(methods_tried))


def extract_with_async_textract(bucket, key):
    start = textract_client.start_document_text_detection(
        DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}}
    )
    job_id = start["JobId"]
    deadline = time.time() + 90  # function timeout is 120s; leave headroom
    status = "IN_PROGRESS"
    while status == "IN_PROGRESS" and time.time() < deadline:
        time.sleep(2)
        result = textract_client.get_document_text_detection(JobId=job_id)
        status = result["JobStatus"]
    if status != "SUCCEEDED":
        raise Exception(f"Textract job status={status} (jobId={job_id})")

    lines, next_token = [], None
    while True:
        if next_token:
            result = textract_client.get_document_text_detection(
                JobId=job_id, NextToken=next_token
            )
        lines += [
            b["Text"] for b in result["Blocks"] if b["BlockType"] == "LINE"
        ]
        next_token = result.get("NextToken")
        if not next_token:
            break
    return "\n".join(lines)


def get_job_requirements(tenant_id, job_id):
    table = dynamodb.Table(JOBS_TABLE)
    resp = table.get_item(Key={"tenant_id": tenant_id, "job_id": job_id})
    if "Item" not in resp:
        raise Exception(f"Job not found: {tenant_id}/{job_id}")
    return resp["Item"]


def score_cv(cv_text, requirements):
    prompt = build_prompt(cv_text, requirements)
    body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 1000,
        "temperature": 0.1,
        "messages": [{"role": "user", "content": prompt}],
    }
    resp = bedrock_client.invoke_model(
        modelId=BEDROCK_MODEL,
        contentType="application/json",
        accept="application/json",
        body=json.dumps(body),
    )
    payload = json.loads(resp["body"].read().decode("utf-8"))
    ai_text = payload["content"][0]["text"]

    try:
        match = re.search(r"\{.*\}", ai_text, re.DOTALL)
        result = json.loads(match.group() if match else ai_text)
        required = ["score", "confidence", "skills_matched",
                    "skills_missing", "years_experience", "summary"]
        for f in required:
            if f not in result:
                raise Exception(f"Missing field: {f}")
        # Coerce numerics defensively.
        result["score"] = int(result["score"])
        result["confidence"] = int(result["confidence"])
        result["years_experience"] = int(result["years_experience"])
        return result
    except Exception as e:
        print(f"Failed to parse AI response: {e} | raw={ai_text[:500]}")
        return {
            "score": 0, "confidence": 0, "skills_matched": [],
            "skills_missing": [], "years_experience": 0,
            "summary": "Could not parse AI scoring response.",
        }


def build_prompt(cv_text, requirements):
    req = requirements.get("skills_required", {})
    pref = requirements.get("skills_preferred", {})
    exp_req = requirements.get("experience_years_required", 2)
    exp_pts = requirements.get("experience_points", 20)
    edu_pts = requirements.get("education_points", 10)
    title = requirements.get("title", "Unknown Role")

    def fmt(d):
        return "\n".join(f"- {k}: {int(v)} points" for k, v in d.items())

    return f"""You are an expert technical recruiter for Fernwood Systems Ltd, scoring a candidate for a {title} role.

SCORING RUBRIC (100 points total):
REQUIRED SKILLS ({int(sum(int(v) for v in req.values()))} pts):
{fmt(req)}
PREFERRED SKILLS ({int(sum(int(v) for v in pref.values()))} pts):
{fmt(pref)}
EXPERIENCE: minimum {exp_req} years ({exp_pts} pts)
EDUCATION: {edu_pts} pts

Score conservatively. Only credit a skill if the CV shows real evidence of it.
"confidence" is your certainty (0-100) in this assessment given CV clarity.

CANDIDATE CV:
{cv_text}

Respond with ONLY valid JSON, no prose, in exactly this shape:
{{"score": 85, "confidence": 90, "skills_matched": ["Docker","AWS"], "skills_missing": ["Kubernetes"], "years_experience": 4, "summary": "One or two sentence assessment."}}"""


def determine_route(score_result, requirements):
    score = score_result["score"]
    confidence = score_result["confidence"]
    shortlist = int(requirements.get("shortlist_threshold", 80))
    review = int(requirements.get("human_review_threshold", 50))
    if confidence < 60:
        return "flagged"
    if score >= shortlist:
        return "shortlisted"
    if score >= review:
        return "under_review"
    return "rejected"


def persist_results(applicant_id, score_result, ai_suggestion, status):
    # Numbers must be Decimal for DynamoDB; lists/strings pass through.
    fields = {
        "score": Decimal(str(score_result["score"])),
        "confidence": Decimal(str(score_result["confidence"])),
        "skills_matched": score_result["skills_matched"],
        "skills_missing": score_result["skills_missing"],
        "years_experience": Decimal(str(score_result["years_experience"])),
        "summary": score_result["summary"],
        "ai_suggested_status": ai_suggestion,
        "status": status,
        "updated_at": _now(),
    }
    expr, names, values = _build_update(fields)
    dynamodb.Table(APPLICATIONS_TABLE).update_item(
        Key={"applicant_id": applicant_id},
        UpdateExpression=expr,
        ExpressionAttributeNames=names,
        ExpressionAttributeValues=values,
    )


def update_status(applicant_id, status):
    expr, names, values = _build_update({"status": status, "updated_at": _now()})
    dynamodb.Table(APPLICATIONS_TABLE).update_item(
        Key={"applicant_id": applicant_id},
        UpdateExpression=expr,
        ExpressionAttributeNames=names,
        ExpressionAttributeValues=values,
    )


def move_to_processed(bucket, key, applicant_id):
    filename = key.split("/")[-1]
    dest = f"processed/{applicant_id}/{filename}"
    s3_client.copy_object(
        CopySource={"Bucket": bucket, "Key": key},
        Bucket=PROCESSED_BUCKET, Key=dest,
    )
    s3_client.delete_object(Bucket=bucket, Key=key)


def trigger_recruiter_alert(applicant_id, score_result):
    """Alert the recruiter about a low-confidence (flagged) application.
    Never emails the candidate — that only happens on a recruiter decision."""
    table = dynamodb.Table(APPLICATIONS_TABLE)
    item = table.get_item(Key={"applicant_id": applicant_id}).get("Item", {})
    payload = {
        "type": "recruiter_alert",
        "applicant_id": applicant_id,
        "candidate_name": item.get("name", "Candidate"),
        "score": score_result["score"],
        "summary": score_result["summary"],
    }
    lambda_client.invoke(
        FunctionName=NOTIFIER_FUNCTION,
        InvocationType="Event",
        Payload=json.dumps(payload),
    )
    print(f"Recruiter alert sent for {applicant_id}")
